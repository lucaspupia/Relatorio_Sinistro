# Imports
import numpy as np
from datetime import datetime
import math
import plotly.express as px
import pandas as pd
import plotly.io as pio
from docx import Document
from docx.shared import Mm
import streamlit as st
from multiprocessing import Value
from tqdm import tqdm

import os
if not os.path.exists("images"):
    os.mkdir("images")

import os
if not os.path.exists("documento"):
    os.mkdir('documento')

global df, data

# cria um novo documento
document = Document()

#Para selecionar a rodovia
st.sidebar.title("Relatório sinistros")
st.sidebar.markdown("Selecione a Rodovia")

rodovias = pd.read_csv('dados/Rodovias.csv')

chart_visual_rodovia = st.sidebar.selectbox('Selecione a Rodovia', rodovias)

Anos = ("2019","2020","2021","2022","2023")
chart_visual_ano = st.sidebar.selectbox('Selecione o ano' ,Anos)

def carregar_csv():
    global df

    filename = 'dados/Acidentes_Bateu_SRE_+_Bateu__19-01-2023.csv'

    df = pd.DataFrame()

    # Inicializar a barra de progresso
    progress_bar = st.progress(0)

    # Open the file and get the total number of rows
    with open(filename, 'r', encoding="utf-8-sig", errors='ignore') as f:
        total_rows = sum(1 for row in f)

    # Use tqdm to display a progress bar
    for chunk in tqdm(pd.read_csv(filename, sep=';', encoding="utf-8-sig", encoding_errors='ignore',
                                  low_memory=False, chunksize=1000), total=total_rows):
        progress_bar.progress(int(chunk.index[-1] / total_rows * 100))
        df = pd.concat([df, chunk])

    # Remover a barra de progresso
    progress_bar.empty()
    # Transformar colunas para seus tipos (astype)
    ## Substituição valores

    df['Tipo da Via'] = np.where(df['Tipo da Via'] == 'Não Informado', 'Z', df['Tipo da Via'])
    df = df.sort_values(by='Tipo da Via', na_position='last')

    df['N° Ocup. Óbito'].replace({"-": 0}, inplace=True)
    df['N° de Ocupantes Feridos'].replace({"-": 0}, inplace=True)

    dateformat = '%Y-%m-%d %H:%M:%S'

    # Convertendo a coluna data de object para datetime
    df['Data Fato'] = pd.to_datetime(df['Data Fato'], format=dateformat,
                                     dayfirst=True, errors='ignore')
    df['data_comunicacao'] = pd.to_datetime(df['data_comunicacao'], format=dateformat,
                                            dayfirst=True)
    df['Validade CNH'] = pd.to_datetime(df['Validade CNH'], format=dateformat,
                                        dayfirst=True, errors='coerce')

    # Transformação para Category

    coluna_categorica = ['Município', 'Bairro', 'Tipo de Acidente', 'Tipo da Via', 'Tipo de Pavimentação da Via',
                         'Conservação da Via', 'Sentido da Via', 'Condição do Semáforo',
                         'Condição da Superfície da Via',
                         'Clima', 'Controle de Tráfego', 'Condição técnica da via', 'Acostamento da Via',
                         'Sinal de Pneu na Pista', 'Sinalização Horizontal', 'Sinalização Vertical',
                         'Sinalização Auxiliar',
                         'Sinalização Obras', 'Sinalização Policial', 'Tipo Pessoa Envolvida',
                         'Nacionalidade Pessoa', 'Sexo Pessoa', 'Estado Civil Pessoa', 'Grau Instrução Pessoa',
                         'Município Pessoa',
                         'Profissão Pessoa', 'Atividade Pessoa', 'Municipio do Local de Trabalho Pessoa', 'MOPP',
                         'Categoria CNH', 'Tempo Habilitação', 'Encaminhado Por', 'Local Encaminhamento',
                         'Município Encaminhamento',
                         'Resultado Vítima', 'Município Veículo', 'Cor do Veículo', 'Marca', 'Modelo',
                         'Categoria do Veículo',
                         'Espécie do Veículo', 'Película no Veículo', 'Veículo no Momento do Fato',
                         'Sinalização Não Identificado',
                         'UF Pessoa', 'UF do Local de Trabalho Pessoa', 'UF do CNH', 'UF do Veículo', 'Unidade']

    df[coluna_categorica] = df[coluna_categorica].astype('category')

    coluna_categorica2 = ['trecho', 'trecho_de', 'trecho_para', 'situacao', 'direita/esquerda', 'sit.',
                          'trechos_coincidentes_a',
                          'trechos_coincidentes_b', 'final_do_trecho_ref']

    df[coluna_categorica2] = df[coluna_categorica2].astype('category')

    # Transformação para Inteiro

    coluna_inteiros = ['Resultado do Acidente - Veículo(s) Danificado(s)', 'Resultado do Acidente - Feridos',
                       'Resultado do Acidente - Óbito(s) no Local', 'Resultado do Acidente - Óbito(s) Posterior',
                       'Resultado do Acidente - PS/Hospital', 'Resultado do Acidente - Não Identificado',
                       'Resultado do Acidente - Outros  ', 'Velocidade Máxima da Via',
                       'Pessoa N°', 'Idade Pessoa', 'Acoplado ao veículo N°', 'Ano Fabricacao Veículo',
                       'N° de Ocupantes no Veículo', 'Nº Km', 'N° de Ocupantes Feridos', 'N° Ocup. Óbito',
                       'Carga do Veículo',
                       'final_do_trecho_alt']
    #######Transformar  Altura da Carga e Peso da Carga, da pra problema em 1,5 #######

    df[coluna_inteiros] = df[coluna_inteiros].astype('float')

    # Transformação para Boleanos

    coluna_bool = ['Produtos Perigosos', 'Acionou Airbag']

    df[coluna_bool] = df[coluna_bool].astype('bool')

    df['Ano Ocorrencia'] = df["Data Fato"].dt.strftime('%Y')
    df['Mês Ocorrencia'] = df["Data Fato"].dt.strftime('%m')
    df['Mês/Ano Ocorrencia'] = df["Data Fato"].dt.strftime('%m/%Y')
    df['Hora Ocorrencia'] = df["Data Fato"].dt.strftime('%H')
    df['dia_da_semana'] = df['dia_da_semana'].astype('str')
    df['semana+hora'] = df[['dia_da_semana', 'Hora Ocorrencia']].agg('-'.join, axis=1)
    df['dia_da_semana'] = df['dia_da_semana'].map({'0': 'Segunda-Feira',
                                                   '1': 'Terça-Feira',
                                                   '2': 'Quarta-Feira',
                                                   '3': 'Quinta-Feira',
                                                   '4': 'Sexta-Feira',
                                                   '5': 'Sábado',
                                                   '6': 'Domingo',
                                                   })

    df['Mês/Ano Ocorrencia'] = pd.to_datetime(df['Mês/Ano Ocorrencia'], format='%m/%Y',
                                              dayfirst=True, errors='coerce')

    df['Tipo de Acidente'] = df['Tipo de Acidente'].str.title()

    df['Tipo de Acidente'] = df['Tipo de Acidente'].map(
        {'Nao Informado': 'Não Informado', 'Nao Identificado': 'Não Informado',
         'Abalroamento Lateral': 'Abalroamento Lateral',
         'Colisao Traseira': 'Colisao Traseira',
         'Abalroamento Transversal': 'Abalroamento Transversal',
         'Colisao Frontal': 'Colisao Frontal',
         'Choque': "Choque",
         'Tombamento': "Tombamento",
         'Acidente Complexo': 'Acidente Complexo',
         "Capotamento": "Capotamento",
         'Queda Moto': "Queda de Moto",
         'Engavetamento': "Engavetamento",
         "Atropelamento": "Atropelamento",
         'Atropelamento Animal': "Atropelamento Animal",
         'Queda Objeto': "Queda de Objeto",
         'Incendio': "Incêndio",
         "Queda Veiculo": "Queda de Veiculo",
         "Queda Passageiro": "Queda de Passageiro",
         "Submerso": "Submerso",
         })

    df["numero"] = df["trecho"].str[:+3]
    df["rod_nome"] = df["trecho"].str[-3:]
    df['Rodovia-'] = df['rod_nome'] + '-' + df['numero']
    df['Rodovia-'] = df['Rodovia'].str.strip('E')

    df['Nº Km'] = df['Nº Km'].round()
carregar_csv()
def sinistros_ano():
    fig = px.bar(data, y='Acidentes', x='Ano Ocorrencia', color='classificacao_acidente', barmode='group',
                 text='Acidentes',
                 category_orders={
                     'classificacao_acidente': ['Sem Vítimas', 'Com Vítimas Feridas', 'Com Vítimas Fatais']},
                 color_discrete_map={"Sem Vítimas": 'green', 'Com Vítimas Feridas': '#ff7f0e',
                                     'Com Vítimas Fatais': '#cc0202'},
                 labels={'classificacao_acidente': 'Classificação do Sinistro',
                         'Nº Km': 'Km'}
                 )

    fig.update_traces(textfont_size=12, textangle=0, textposition="inside", cliponaxis=False)

    fig.update_layout(bargap=0.6)

    fig.update_traces(width=0.4)

    fig.for_each_annotation(lambda a: a.update(text=a.text.split("=")[1]))

    fig.update_yaxes(
        title_text="Sinistros",
        title_standoff=25,
        dtick=5,
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        range=[0, data['Acidentes'].max() + 5])

    fig.update_xaxes(
        title_standoff=25,
        dtick=1,
        # griddash='dash',
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        #         range=[data['Nº Km'].min()-0.5, data['Nº Km'].max()+0.5]
    )

    # Update plot sizing
    fig.update_layout(
        width=1000,
        height=500,
        autosize=True,
        margin=dict(t=100, b=100, l=100, r=250),
        title="Sinistros por Ano (2019 a 2022)<br>Rodovia: " + Rodovia
    )

    fig.update_layout(legend=dict(
        orientation="v",
        y=1.02,
        x=1.02
    ))

    fig.update_layout(paper_bgcolor="rgba(230, 230, 230, 0.0)", plot_bgcolor="rgba(230, 230, 230, 0.3)")

    hoje = datetime.today().strftime('%d-%m-%Y')

    caminho = '1_' + str(Trecho) + '_ano_' + str(hoje) + '.png'

    pio.write_image(fig, "images/" + caminho)

    document.add_picture("images/" + caminho, width=Mm(16 * 10), height=Mm(8 * 10))

    st.plotly_chart(fig)
    #st.sidebar.markdown('##Ano')

def semana_hora():
    fig = px.density_heatmap(data, x="dia_da_semana", y="Hora Ocorrencia", z='Acidentes',
                             text_auto=True,
                             category_orders={
                                 'Hora Ocorrencia': ['00', '01', '02', '03', '04', '05', '06', '07', '08', '09', '10',
                                                     '11', '12', '13', '14', '15', '16', '17', '18', '19', '20', '21',
                                                     '22', '23'],
                                 'dia_da_semana': ['Segunda-Feira', 'Terça-Feira', 'Quarta-Feira', 'Quinta-Feira',
                                                   'Sexta-Feira', 'Sábado', 'Domingo'],
                                 'classificacao_acidente': ['Sem Vítimas', 'Com Vítimas Feridas',
                                                            'Com Vítimas Fatais']},

                             color_continuous_scale='greens',
                             labels={'classificacao_acidente': 'Classificação do Sinistro',
                                     'dia_da_semana': 'Dia da Semana',
                                     'Hora Ocorrencia': 'Hora da Ocorrência',
                                     'Acidentes': 'Sinistros'})

    # fig.update_traces(texttemplate='%{text:.2s}', textposition='outside')

    fig.update_layout(uniformtext_minsize=8, uniformtext_mode='hide')

    fig.for_each_annotation(lambda a: a.update(text=a.text.split("=")[1]))

    fig.update_yaxes(
        #        title_text = "Sinistros",
        title_standoff=25,
        dtick=1,
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        range=[-0.5, 23.5]
    )

    fig.update_xaxes(
        title_standoff=25,
        dtick=1,
        #        griddash='dash',
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        #        range=[data['Acidentes'].min()-0.2, data['Acidentes'].max()+0.2]
    )

    # Update plot sizing
    fig.update_layout(
        width=500,
        height=500,
        autosize=True,
        margin=dict(t=100, b=100, l=0, r=0),
        title="Sinistros classificados por Dia da Semana e Hora<br>Rodovia: " + Rodovia
    )

    fig.update_layout(paper_bgcolor="rgba(230, 230, 230, 0.0)", plot_bgcolor="rgba(230, 230, 230, 0.3)")

    hoje = datetime.today().strftime('%d-%m-%Y')

    caminho = '5_' + str(Trecho) + '_semana_hora_' + str(hoje) + '.png'
    pio.write_image(fig, "images/" + caminho)

    document.add_picture("images/" + caminho, width=Mm(10 * 10), height=Mm(10 * 10))
    st.plotly_chart(fig)
def semana_hora_year_class():
    fig = px.density_heatmap(data, x="dia_da_semana", y="Hora Ocorrencia", z='Acidentes',
                             text_auto=True,
                             facet_row='Ano Ocorrencia', facet_col='classificacao_acidente',
                             category_orders={
                                 'Hora Ocorrencia': ['00', '01', '02', '03', '04', '05', '06', '07', '08', '09', '10',
                                                     '11', '12', '13', '14', '15', '16', '17', '18', '19', '20', '21',
                                                     '22', '23'],
                                 'dia_da_semana': ['Segunda-Feira', 'Terça-Feira', 'Quarta-Feira', 'Quinta-Feira',
                                                   'Sexta-Feira', 'Sábado', 'Domingo'],
                                 'classificacao_acidente': ['Sem Vítimas', 'Com Vítimas Feridas',
                                                            'Com Vítimas Fatais']},

                             color_continuous_scale='Greens',
                             labels={'classificacao_acidente': 'Classificação do Sinistro',
                                     'dia_da_semana': 'Dia da Semana',
                                     'Hora Ocorrencia': 'Hora da Ocorrência',
                                     'Acidentes': 'Sinistros'}
                             )

    # fig.update_traces(texttemplate='%{text:.2s}', textposition='outside')

    fig.update_layout(uniformtext_minsize=8, uniformtext_mode='hide')

    fig.for_each_annotation(lambda a: a.update(text=a.text.split("=")[1]))

    fig.update_yaxes(
        #        title_text = "Sinistros",
        title_standoff=25,
        dtick=1,
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        range=[-0.5, 23.5]
    )

    fig.update_xaxes(
        title_standoff=25,
        dtick=1,
        #        griddash='dash',
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        #        range=[data['Acidentes'].min()-0.2, data['Acidentes'].max()+0.2]
    )

    # Update plot sizing
    fig.update_layout(
        width=1100,
        height=1300,
        autosize=True,
        margin=dict(t=110, b=0, l=0, r=0),
        title="Sinistros classificados por Dia da Semana, Hora, Ano e Gravidade<br>Rodovia: " + Rodovia
    )

    fig.update_layout(paper_bgcolor="rgba(230, 230, 230, 0.0)", plot_bgcolor="rgba(230, 230, 230, 0.3)")

    hoje = datetime.today().strftime('%d-%m-%Y')

    caminho = '2_' + str(Trecho) + '_semana_hora_year_class_' + str(hoje) + '.png'

    pio.write_image(fig, "images/" + caminho)

    document.add_picture("images/" + caminho, height=Mm(18 * 10), width=Mm(15.23 * 10))

    st.plotly_chart(fig)
def semana_hora_year():
    fig = px.density_heatmap(data, x="dia_da_semana", y="Hora Ocorrencia", z='Acidentes',
                             text_auto=True,
                             facet_row='Ano Ocorrencia',
                             category_orders={
                                 'Hora Ocorrencia': ['00', '01', '02', '03', '04', '05', '06', '07', '08', '09', '10',
                                                     '11', '12', '13', '14', '15', '16', '17', '18', '19', '20', '21',
                                                     '22', '23'],
                                 'dia_da_semana': ['Segunda-Feira', 'Terça-Feira', 'Quarta-Feira', 'Quinta-Feira',
                                                   'Sexta-Feira', 'Sábado', 'Domingo'],
                                 'classificacao_acidente': ['Sem Vítimas', 'Com Vítimas Feridas',
                                                            'Com Vítimas Fatais']},

                             color_continuous_scale='Greens',
                             labels={'Acidentes': 'Sinistros',
                                     'dia_da_semana': 'Dia da Semana',
                                     'Hora Ocorrencia': 'Hora da Ocorrência'}
                             )

    # fig.update_traces(texttemplate='%{text:.2s}', textposition='outside')

    fig.update_layout(uniformtext_minsize=8, uniformtext_mode='hide')

    fig.for_each_annotation(lambda a: a.update(text=a.text.split("=")[1]))

    fig.update_yaxes(
        #        title_text = "Sinistros",
        title_standoff=25,
        dtick=1,
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        range=[-0.5, 23.5]
    )

    fig.update_xaxes(
        title_standoff=25,
        dtick=1,
        #        griddash='dash',
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        #        range=[data['Acidentes'].min()-0.2, data['Acidentes'].max()+0.2]
    )

    # Update plot sizing
    fig.update_layout(
        width=500,
        height=1300,
        autosize=True,
        margin=dict(t=110, b=0, l=0, r=0),
        title="Sinistros classificados por Dia da Semana, Hora e Ano<br>Rodovia: " + Rodovia
    )

    fig.update_layout(paper_bgcolor="rgba(230, 230, 230, 0.0)", plot_bgcolor="rgba(230, 230, 230, 0.3)")

    hoje = datetime.today().strftime('%d-%m-%Y')

    caminho = '3_' + str(Trecho) + '_semana_hora_year_' + str(hoje) + '.png'
    pio.write_image(fig, "images/" + caminho)

    document.add_picture("images/" + caminho, height=Mm(23.4 * 10), width=Mm(9 * 10))

    st.plotly_chart(fig)
def semana_hora_class():
    fig = px.density_heatmap(data, x="dia_da_semana", y="Hora Ocorrencia", z='Acidentes',
                             text_auto=True,
                             facet_col='classificacao_acidente',
                             category_orders={
                                 'Hora Ocorrencia': ['00', '01', '02', '03', '04', '05', '06', '07', '08', '09', '10',
                                                     '11', '12', '13', '14', '15', '16', '17', '18', '19', '20', '21',
                                                     '22', '23'],
                                 'dia_da_semana': ['Segunda-Feira', 'Terça-Feira', 'Quarta-Feira', 'Quinta-Feira',
                                                   'Sexta-Feira', 'Sábado', 'Domingo'],
                                 'classificacao_acidente': ['Sem Vítimas', 'Com Vítimas Feridas',
                                                            'Com Vítimas Fatais']},

                             color_continuous_scale='Greens',
                             labels={'classificacao_acidente': 'Classificação do Sinistro',
                                     'dia_da_semana': 'Dia da Semana',
                                     'Hora Ocorrencia': 'Hora da Ocorrência',
                                     'Acidentes': 'Sinistros'}
                             )

    # fig.update_traces(texttemplate='%{text:.2s}', textposition='outside')

    fig.update_layout(uniformtext_minsize=8, uniformtext_mode='hide')

    fig.for_each_annotation(lambda a: a.update(text=a.text.split("=")[1]))

    fig.update_yaxes(
        #        title_text = "Sinistros",
        title_standoff=25,
        dtick=1,
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        range=[-0.5, 23.5]
    )

    fig.update_xaxes(
        title_standoff=25,
        dtick=1,
        #        griddash='dash',
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        #        range=[data['Acidentes'].min()-0.2, data['Acidentes'].max()+0.2]
    )

    # Update plot sizing
    fig.update_layout(
        width=1100,
        height=500,
        autosize=True,
        margin=dict(t=110, b=0, l=0, r=0),
        title="Sinistros classificados por Dia da Semana, Hora e Gravidade<br>Rodovia: " + Rodovia
    )

    fig.update_layout(paper_bgcolor="rgba(230, 230, 230, 0.0)", plot_bgcolor="rgba(230, 230, 230, 0.3)")

    hoje = datetime.today().strftime('%d-%m-%Y')

    caminho = '4_' + str(Trecho) + '_semana_hora_class_' + str(hoje) + '.png'
    pio.write_image(fig, "images/" + caminho)

    document.add_picture("images/" + caminho, height=Mm(7.27 * 10), width=Mm(16 * 10))

    st.plotly_chart(fig)
def veiculo_year():
    fig = px.bar(df_veiculos1, y='Tipo Veículo_x', x='Acidentes', color='classificacao_acidente',
                 facet_row='Ano Ocorrencia', text='Acidentes',
                 category_orders={
                     'Tipo Veículo': ['AUTOMÓVEL', 'CAMINHÃO', 'MOTOCICLETA', 'UTILITÁRIO', 'ONIBUS', 'BICICLETA',
                                      'OUTROS'],
                     'classificacao_acidente': ['Sem Vítimas', 'Com Vítimas Feridas', 'Com Vítimas Fatais']},
                 labels={'classificacao_acidente': 'Classificação do Sinistro',
                         'Tipo Veículo_x': "Tipo de Veículo",
                         'Acidentes': 'Sinistros'},
                 color_discrete_map={"Sem Vítimas": 'green', 'Com Vítimas Feridas': '#ff7f0e',
                                     'Com Vítimas Fatais': '#cc0202'},
                 )

    fig.update_traces(textfont_size=12, textangle=0)

    fig.update_layout(bargap=0)

    fig.update_traces(width=0.8)

    fig.for_each_annotation(lambda a: a.update(text=a.text.split("=")[1]))

    fig.update_yaxes(
        title_text="Tipo de Veículos",
        title_standoff=25,
        dtick=1,
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        #        range=[0, data['Acidentes'].max()+2]
    )

    fig.update_xaxes(
        title_standoff=25,
        #        dtick=1,
        #        griddash='dash',
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        #        range=[data['Acidentes'].min()-0.2, data['Acidentes'].max()+0.2]
    )

    # Update plot sizing
    fig.update_layout(
        width=1000,
        height=1000,
        autosize=True,
        margin=dict(t=100, b=0, l=100, r=200),
        title="Sinistros classificados por Tipo de Veiculo por Ano<br>Rodovia: " + Rodovia
    )

    fig.update_layout(legend=dict(
        orientation="v",
        y=1.01,
        x=1.02
    ))

    # fig.update_layout(xaxis=dict(rangeslider=dict(visible=True),rangeslider_thickness = 0.02))

    fig.update_layout(paper_bgcolor="rgba(230, 230, 230, 0.0)", plot_bgcolor="rgba(230, 230, 230, 0.3)")

    hoje = datetime.today().strftime('%d-%m-%Y')

    caminho = '6_' + str(Trecho) + '_veiculo_year_' + str(hoje) + '.png'

    pio.write_image(fig, "images/" + caminho)

    document.add_picture("images/" + caminho, height=Mm(16 * 10), width=Mm(16 * 10))

    st.plotly_chart(fig)
def veiculo():
    df_veiculos2 = df_veiculos.groupby(['Tipo Veículo_x', 'Rodovia', 'classificacao_acidente'])[
        'Protocolo BATEU'].count().reset_index().rename(columns={'Protocolo BATEU': 'Acidentes'})

    fig = px.bar(df_veiculos2, y='Tipo Veículo_x', x='Acidentes', color='classificacao_acidente',
                 text='Acidentes',
                 category_orders={
                     'Tipo Veículo': ['AUTOMÓVEL', 'CAMINHÃO', 'MOTOCICLETA', 'UTILITÁRIO', 'ONIBUS', 'BICICLETA',
                                      'OUTROS'],
                     'classificacao_acidente': ['Sem Vítimas', 'Com Vítimas Feridas', 'Com Vítimas Fatais']},
                 labels={'classificacao_acidente': 'Classificação do Sinistro',
                         'Acidentes': 'Sinistros'},
                 color_discrete_map={"Sem Vítimas": 'green', 'Com Vítimas Feridas': '#ff7f0e',
                                     'Com Vítimas Fatais': '#cc0202'})

    # fig.update_traces(textposition='inside')

    fig.update_layout(bargap=0)

    fig.update_traces(width=0.8)

    fig.for_each_annotation(lambda a: a.update(text=a.text.split("=")[1]))

    fig.update_yaxes(
        title_text="Tipo dos Veículos",
        title_standoff=25,
        dtick=1,
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        #        range=[0, data['Acidentes'].max()+2]
    )

    fig.update_xaxes(
        title_standoff=25,
        #        dtick=5,
        #        griddash='dash',
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        #        range=[data['Acidentes'].min()-0.2, data['Acidentes'].max()+0.2]
    )

    # Update plot sizing
    fig.update_layout(
        width=1000,
        height=400,
        autosize=True,
        margin=dict(t=100, b=0, l=100, r=200),
        title="Sinistros classificados por Tipo de Veiculo (2019 a 2022)<br>Rodovia: " + Rodovia
    )

    fig.update_layout(legend=dict(
        orientation="v",
        y=1.04,
        x=1.02
    ))

    # fig.update_layout(xaxis=dict(rangeslider=dict(visible=True),rangeslider_thickness = 0.02))

    fig.update_layout(paper_bgcolor="rgba(230, 230, 230, 0.0)", plot_bgcolor="rgba(230, 230, 230, 0.3)")

    hoje = datetime.today().strftime('%d-%m-%Y')

    caminho = '7_' + str(Trecho) + '_veiculo_' + str(hoje) + '.png'
    pio.write_image(fig, "images/" + caminho)

    document.add_picture("images/" + caminho, height=Mm(6.4 * 10), width=Mm(16 * 10))

    st.plotly_chart(fig)
def km_year():
    data = df_acidentes_ano.groupby(['Ano Ocorrencia', 'Nº Km', 'classificacao_acidente', 'Rodovia'])[
        'Protocolo BATEU'].count().reset_index().rename(columns={'Protocolo BATEU': 'Acidentes'})

    fig = px.bar(data, y='Acidentes', x='Nº Km', color='classificacao_acidente', barmode='group',
                 facet_row='Ano Ocorrencia', text='Acidentes',
                 category_orders={
                     'classificacao_acidente': ['Sem Vítimas', 'Com Vítimas Feridas', 'Com Vítimas Fatais']},
                 color_discrete_map={"Sem Vítimas": 'green', 'Com Vítimas Feridas': '#ff7f0e',
                                     'Com Vítimas Fatais': '#cc0202'},
                 labels={'classificacao_acidente': 'Classificação do Sinistro',
                         'Nº Km': 'Km'}
                 )

    fig.update_traces(textfont_size=12, textangle=0, textposition="outside", cliponaxis=False)

    fig.update_layout(bargap=0.6)

    fig.update_traces(width=0.2)

    fig.for_each_annotation(lambda a: a.update(text=a.text.split("=")[1]))

    fig.update_yaxes(
        title_text="Sinistros",
        title_standoff=25,
        dtick=2,
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        range=[0, data['Acidentes'].max() + 2]
    )

    fig.update_xaxes(
        title_standoff=25,
        dtick=1,
        #        griddash='dash',
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        range=[data['Nº Km'].min() - 0.2, data['Nº Km'].max() + 0.2])

    # Update plot sizing
    fig.update_layout(
        width=1000,
        height=1000,
        autosize=True,
        margin=dict(t=100, b=0, l=100, r=200),
        title="Sinistros classificados por Ano x Km<br>Rodovia: " + Rodovia
    )

    fig.update_layout(legend=dict(
        orientation="v",
        y=1.01,
        x=1.02
    ))

    # fig.update_layout(xaxis=dict(rangeslider=dict(visible=True),rangeslider_thickness = 0.02))

    fig.update_layout(paper_bgcolor="rgba(230, 230, 230, 0.0)", plot_bgcolor="rgba(230, 230, 230, 0.3)")

    hoje = datetime.today().strftime('%d-%m-%Y')

    caminho = '8_' + str(Trecho) + '_km_year_' + str(hoje) + '.png'
    pio.write_image(fig, "images/" + caminho)

    document.add_picture("images/" + caminho, height=Mm(16 * 10), width=Mm(16 * 10))

    st.plotly_chart(fig)
def km():
    data = df_acidentes_ano.groupby(['Nº Km', 'classificacao_acidente', 'Rodovia'])[
        'Protocolo BATEU'].count().reset_index().rename(columns={'Protocolo BATEU': 'Acidentes'})

    fig = px.bar(data, y='Acidentes', x='Nº Km', color='classificacao_acidente', barmode='group',
                 text='Acidentes',
                 category_orders={
                     'classificacao_acidente': ['Sem Vítimas', 'Com Vítimas Feridas', 'Com Vítimas Fatais']},
                 color_discrete_map={"Sem Vítimas": 'green', 'Com Vítimas Feridas': '#ff7f0e',
                                     'Com Vítimas Fatais': '#cc0202'},
                 labels={'classificacao_acidente': 'Classificação do Sinistro',
                         'Nº Km': 'Km'}
                 )

    fig.update_traces(textfont_size=12, textangle=0, textposition="inside", cliponaxis=False)

    fig.update_layout(bargap=0.6)

    fig.update_traces(width=0.4)

    fig.for_each_annotation(lambda a: a.update(text=a.text.split("=")[1]))

    fig.update_yaxes(
        title_text="Sinistros",
        title_standoff=25,
        dtick=2,
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        range=[0, data['Acidentes'].max() + 2])

    fig.update_xaxes(
        title_standoff=25,
        dtick=1,
        #        griddash='dash',
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        range=[data['Nº Km'].min() - 0.5, data['Nº Km'].max() + 0.5])

    # Update plot sizing
    fig.update_layout(
        width=1000,
        height=500,
        autosize=True,
        margin=dict(t=100, b=100, l=100, r=250),
        title="Sinistros por Km (2019 a 2022)<br>Rodovia: " + Rodovia
    )

    fig.update_layout(legend=dict(
        orientation="v",
        y=1.02,
        x=1.02
    ))

    fig.update_layout(paper_bgcolor="rgba(230, 230, 230, 0.0)", plot_bgcolor="rgba(230, 230, 230, 0.3)")

    hoje = datetime.today().strftime('%d-%m-%Y')

    caminho = '9_' + str(Trecho) + '_km_' + str(hoje) + '.png'
    pio.write_image(fig, "images/" + caminho)

    document.add_picture("images/" + caminho, height=Mm(6 * 10), width=Mm(12 * 10))

    st.plotly_chart(fig)
def mes_year():
    data = df_acidentes_ano.groupby(['Ano Ocorrencia', 'Mês Ocorrencia', 'classificacao_acidente', 'Rodovia'])[
        'Protocolo BATEU'].count().reset_index().rename(columns={'Protocolo BATEU': 'Acidentes'})

    fig = px.bar(data, y='Acidentes', x='Mês Ocorrencia', color='classificacao_acidente', barmode='group',
                 facet_row='Ano Ocorrencia', text='Acidentes',
                 category_orders={
                     'classificacao_acidente': ['Sem Vítimas', 'Com Vítimas Feridas', 'Com Vítimas Fatais'],
                     'Mês Ocorrencia': ['01', '02', '03', '04', '05', '06', '07', '08', '09', '10', '11', '12']},
                 color_discrete_map={"Sem Vítimas": 'green', 'Com Vítimas Feridas': '#ff7f0e',
                                     'Com Vítimas Fatais': '#cc0202'},
                 labels={'classificacao_acidente': 'Classificação do Sinistro',
                         'Mês Ocorrencia': 'Mês'})

    fig.update_traces(textfont_size=12, textangle=0, textposition="outside", cliponaxis=False)

    fig.update_layout(bargap=0.6)

    fig.update_traces(width=0.2)

    fig.for_each_annotation(lambda a: a.update(text=a.text.split("=")[1]))

    fig.update_yaxes(
        title_text="Sinistros",
        title_standoff=25,
        dtick=2,
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        range=[0, data['Acidentes'].max() + 5])

    fig.update_xaxes(
        title_standoff=25,
        dtick=1,
        #        griddash='dash',
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        range=[-0.5, 11.5]
    )

    # Update plot sizing
    fig.update_layout(
        width=1000,
        height=1000,
        autosize=True,
        margin=dict(t=100, b=0, l=100, r=200),
        title="Sinistros classificados por Ano x Mês<br>Rodovia: " + Rodovia
    )

    fig.update_layout(legend=dict(
        orientation="v",
        y=1.01,
        x=1.02
    ))

    fig.update_layout(paper_bgcolor="rgba(230, 230, 230, 0.0)", plot_bgcolor="rgba(230, 230, 230, 0.3)")

    hoje = datetime.today().strftime('%d-%m-%Y')

    caminho = '10_' + str(Trecho) + '_mes_year_' + str(hoje) + '.png'
    pio.write_image(fig, "images/" + caminho)

    document.add_picture("images/" + caminho, height=Mm(15 * 10), width=Mm(15 * 10))

    st.plotly_chart(fig)
def mes():
    data = df_acidentes_ano.groupby(['Mês Ocorrencia', 'classificacao_acidente', 'Rodovia'])[
        'Protocolo BATEU'].count().reset_index().rename(columns={'Protocolo BATEU': 'Acidentes'})

    fig = px.bar(data, y='Acidentes', x='Mês Ocorrencia', color='classificacao_acidente', barmode='group',
                 text='Acidentes',
                 category_orders={
                     'classificacao_acidente': ['Sem Vítimas', 'Com Vítimas Feridas', 'Com Vítimas Fatais'],
                     'Mês Ocorrencia': ['01', '02', '03', '04', '05', '06', '07', '08', '09', '10', '11', '12']},
                 color_discrete_map={"Sem Vítimas": 'green', 'Com Vítimas Feridas': '#ff7f0e',
                                     'Com Vítimas Fatais': '#cc0202'},
                 labels={'Mês Ocorrencia': "Mês",
                         'classificacao_acidente': 'Classificação do Sinistro'})

    fig.update_traces(textfont_size=12, textangle=0, textposition="inside", cliponaxis=False)

    fig.update_layout(bargap=0.6)

    fig.update_traces(width=0.4)

    fig.for_each_annotation(lambda a: a.update(text=a.text.split("=")[1]))

    fig.update_yaxes(
        title_text="Sinistros",
        title_standoff=25,
        dtick=5,
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        range=[0, data['Acidentes'].max() + 5])

    fig.update_xaxes(
        title_standoff=25,
        dtick=1,
        #        griddash='dash',
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        range=[-0.5, 11.5]
    )

    # Update plot sizing
    fig.update_layout(
        width=1000,
        height=500,
        autosize=True,
        margin=dict(t=100, b=100, l=100, r=250),
        title="Sinistros por Mês (2019 a 2022)<br>Rodovia: " + Rodovia
    )

    fig.update_layout(legend=dict(
        orientation="v",
        y=1.02,
        x=1.02
    ))

    fig.update_layout(paper_bgcolor="rgba(230, 230, 230, 0.0)", plot_bgcolor="rgba(230, 230, 230, 0.3)")

    hoje = datetime.today().strftime('%d-%m-%Y')

    caminho = '11_' + str(Trecho) + '_mes_' + str(hoje) + '.png'
    pio.write_image(fig, "images/" + caminho)

    document.add_picture("images/" + caminho, height=Mm(6.5 * 10), width=Mm(13 * 10))

    st.plotly_chart(fig)
def tipo_year():
    data = df_acidentes_ano.groupby(['Ano Ocorrencia', 'Tipo de Acidente', 'classificacao_acidente', 'Rodovia'])[
        'Protocolo BATEU'].count().reset_index().rename(columns={'Protocolo BATEU': 'Acidentes'})

    fig = px.bar(data, y='Acidentes', x='Tipo de Acidente', color='classificacao_acidente', barmode='group',
                 facet_row='Ano Ocorrencia', text='Acidentes',
                 category_orders={
                     'classificacao_acidente': ['Sem Vítimas', 'Com Vítimas Feridas', 'Com Vítimas Fatais']},
                 color_discrete_map={"Sem Vítimas": 'green', 'Com Vítimas Feridas': '#ff7f0e',
                                     'Com Vítimas Fatais': '#cc0202'},
                 labels={'classificacao_acidente': 'Classificação do Sinistro',
                         'Tipo de Acidente': ''}
                 )

    fig.update_traces(textfont_size=12, textangle=0, textposition="outside", cliponaxis=False)

    fig.update_layout(bargap=0.6)

    fig.update_traces(width=0.2)

    fig.for_each_annotation(lambda a: a.update(text=a.text.split("=")[1]))

    fig.update_yaxes(
        title_text="Sinistros",
        title_standoff=25,
        dtick=5,
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        range=[0, data['Acidentes'].max() + 5])

    fig.update_xaxes(
        title_standoff=25,
        dtick=1,
        #        griddash='dash',
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        #        range=[data['Tipo de Acidente'].min()-0.5, data['Mês Ocorrencia'].max()+0.5]
    )

    # Update plot sizing
    fig.update_layout(
        width=1200,
        height=1500,
        autosize=True,
        margin=dict(t=100, b=0, l=100, r=200),
        title="Sinistros classificados por Ano e Tipo<br>Rodovia: " + Rodovia
    )

    fig.update_layout(legend=dict(
        orientation="v",
        y=1.005,
        x=1.02
    ))

    fig.update_layout(paper_bgcolor="rgba(230, 230, 230, 0.0)", plot_bgcolor="rgba(230, 230, 230, 0.3)")

    hoje = datetime.today().strftime('%d-%m-%Y')

    caminho = '12_' + str(Trecho) + '_tipo_year_' + str(hoje) + '.png'
    pio.write_image(fig, "images/" + caminho)

    document.add_picture("images/" + caminho, height=Mm(16 * 10), width=Mm(12.8 * 10))

    st.plotly_chart(fig)
def tipo():
    data = df_acidentes_ano.groupby(['Tipo de Acidente', 'classificacao_acidente', 'Rodovia'])[
        'Protocolo BATEU'].count().reset_index().rename(columns={'Protocolo BATEU': 'Acidentes'})

    fig = px.bar(data, y='Acidentes', x='Tipo de Acidente', color='classificacao_acidente', barmode='group',
                 text='Acidentes',
                 category_orders={
                     'classificacao_acidente': ['Sem Vítimas', 'Com Vítimas Feridas', 'Com Vítimas Fatais']},
                 color_discrete_map={"Sem Vítimas": 'green', 'Com Vítimas Feridas': '#ff7f0e',
                                     'Com Vítimas Fatais': '#cc0202'},
                 labels={'classificacao_acidente': 'Classificação do Sinistro',
                         'Tipo de Acidente': ''}
                 )

    fig.update_traces(textfont_size=12, textangle=0, textposition="outside", cliponaxis=False)

    fig.update_layout(bargap=0.6)

    fig.update_traces(width=0.4)

    fig.for_each_annotation(lambda a: a.update(text=a.text.split("=")[1]))

    fig.update_yaxes(
        title_text="Sinistros",
        title_standoff=25,
        dtick=10,
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        range=[0, data['Acidentes'].max() + 5])

    fig.update_xaxes(
        title_standoff=25,
        dtick=1,
        #        griddash='dash',
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        #        range=[data['Mês Ocorrencia'].min()-0.5, data['Nº Km'].max()+0.5]
    )

    # Update plot sizing
    fig.update_layout(
        width=1000,
        height=500,
        autosize=True,
        margin=dict(t=100, b=100, l=100, r=250),
        title="Sinistros por Tipo (2019 a 2022)<br>Rodovia: " + Rodovia
    )

    fig.update_layout(legend=dict(
        orientation="v",
        y=1.02,
        x=1.02
    ))

    fig.update_layout(paper_bgcolor="rgba(230, 230, 230, 0.0)", plot_bgcolor="rgba(230, 230, 230, 0.3)")

    hoje = datetime.today().strftime('%d-%m-%Y')

    caminho = '13_' + str(Trecho) + '_tipo_' + str(hoje) + '.png'
    pio.write_image(fig, "images/" + caminho)

    document.add_picture("images/" + caminho, height=Mm(7.4 * 10), width=Mm(14.8 * 10))

    st.plotly_chart(fig)
def clima_year():
    data = df_acidentes_ano.groupby(['Ano Ocorrencia', 'Clima', 'classificacao_acidente', 'Rodovia'])[
        'Protocolo BATEU'].count().reset_index().rename(columns={'Protocolo BATEU': 'Acidentes'})

    fig = px.bar(data, y='Acidentes', x='Clima', color='classificacao_acidente', barmode='group',
                 facet_row='Ano Ocorrencia', text='Acidentes',
                 category_orders={
                     'classificacao_acidente': ['Sem Vítimas', 'Com Vítimas Feridas', 'Com Vítimas Fatais']},
                 color_discrete_map={"Sem Vítimas": 'green', 'Com Vítimas Feridas': '#ff7f0e',
                                     'Com Vítimas Fatais': '#cc0202'},
                 labels={'classificacao_acidente': 'Classificação do Sinistro',
                         'Clima': ''}
                 )

    fig.update_traces(textfont_size=12, textangle=0, textposition="auto", cliponaxis=False)

    fig.update_layout(bargap=0.6)

    fig.update_traces(width=0.2)

    fig.for_each_annotation(lambda a: a.update(text=a.text.split("=")[1]))

    fig.update_yaxes(
        title_text="Sinistros",
        title_standoff=25,
        dtick=10,
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        range=[0, data['Acidentes'].max() + 10])

    fig.update_xaxes(
        title_standoff=25,
        dtick=1,
        #        griddash='dash',
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        #        range=[data['Tipo de Acidente'].min()-0.5, data['Mês Ocorrencia'].max()+0.5]
    )

    # Update plot sizing
    fig.update_layout(
        width=1200,
        height=1100,
        autosize=True,
        margin=dict(t=80, b=50, l=100, r=200),
        title="Sinistros classificados por Situação Climática por Ano<br>Rodovia: " + Rodovia
    )

    fig.update_layout(legend=dict(
        orientation="v",
        y=1.01,
        x=1.02
    ))

    fig.update_layout(paper_bgcolor="rgba(230, 230, 230, 0.0)", plot_bgcolor="rgba(230, 230, 230, 0.3)")

    hoje = datetime.today().strftime('%d-%m-%Y')

    caminho = '14_' + str(Trecho) + '_clima_year_' + str(hoje) + '.png'
    pio.write_image(fig, "images/" + caminho)

    document.add_picture("images/" + caminho, height=Mm(15 * 10), width=Mm(16.36 * 10))

    st.plotly_chart(fig)
def clima():
    data = df_acidentes_ano.groupby(['Clima', 'classificacao_acidente', 'Rodovia'])[
        'Protocolo BATEU'].count().reset_index().rename(columns={'Protocolo BATEU': 'Acidentes'})

    fig = px.bar(data, y='Acidentes', x='Clima', color='classificacao_acidente', barmode='group',
                 text='Acidentes',
                 category_orders={
                     'classificacao_acidente': ['Sem Vítimas', 'Com Vítimas Feridas', 'Com Vítimas Fatais']},
                 color_discrete_map={"Sem Vítimas": 'green', 'Com Vítimas Feridas': '#ff7f0e',
                                     'Com Vítimas Fatais': '#cc0202'},
                 labels={'classificacao_acidente': 'Classificação do Sinistro',
                         'Clima': ''}
                 )

    fig.update_traces(textfont_size=12, textangle=0, textposition="auto", cliponaxis=False)

    fig.update_layout(bargap=0.6)

    fig.update_traces(width=0.4)

    fig.for_each_annotation(lambda a: a.update(text=a.text.split("=")[1]))

    fig.update_yaxes(
        title_text="Sinistros",
        title_standoff=25,
        dtick=5,
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        range=[0, data['Acidentes'].max() + 10])

    fig.update_xaxes(
        title_standoff=25,
        dtick=1,
        #        griddash='dash',
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        #        range=[data['Mês Ocorrencia'].min()-0.5, data['Nº Km'].max()+0.5]
    )

    # Update plot sizing
    fig.update_layout(
        width=1000,
        height=500,
        autosize=True,
        margin=dict(t=100, b=100, l=100, r=250),
        title="Sinistros por Situação Climática (2019 a 2022)<br>Rodovia: " + Rodovia
    )

    fig.update_layout(legend=dict(
        orientation="v",
        y=1.02,
        x=1.02
    ))

    fig.update_layout(paper_bgcolor="rgba(230, 230, 230, 0.0)", plot_bgcolor="rgba(230, 230, 230, 0.3)")

    hoje = datetime.today().strftime('%d-%m-%Y')
    caminho = '15_' + str(Trecho) + '_clima_' + str(hoje) + '.png'
    pio.write_image(fig, "images/" + caminho)

    document.add_picture("images/" + caminho, height=Mm(7.4 * 10), width=Mm(14.8 * 10))

    st.plotly_chart(fig)
def box_sex_year():
    fig = px.box(df_pessoas2, x='Tipo Pessoa Envolvida', y='Idade Pessoa',
                 color="Sexo Pessoa", facet_col='classificacao_acidente',
                 color_discrete_map={"M": '#050df7', 'F': '#f70505'},
                 category_orders={'Tipo Pessoa Envolvida': ['Condutor', 'Passageiro'],
                                  'Sexo Pessoa': ['M', 'F'],
                                  'Ano Ocorrencia': ['2019', '2020', '2021', '2022'],
                                  'classificacao_acidente': ['Sem Vítimas', 'Com Vítimas Feridas',
                                                             'Com Vítimas Fatais']},
                 points='all',
                 labels={'Sexo Pessoa': 'Gênero',
                         'Tipo Pessoa Envolvida': '',
                         'Idade Pessoa': 'Idade'}
                 )

    fig.update_layout(uniformtext_minsize=8, uniformtext_mode='hide')

    fig.for_each_annotation(lambda a: a.update(text=a.text.split("=")[1]))

    fig.update_yaxes(
        #        title_text = "Sinistros",
        title_standoff=25,
        dtick=10,
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        range=[0, df_pessoas2['Idade Pessoa'].max() + 20])

    fig.update_xaxes(
        title_standoff=25,
        dtick=1,
        #        griddash='dash',
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',

        #        range=[data['Mês Ocorrencia'].min()-0.5, data['Nº Km'].max()+0.5]
    )

    # Update plot sizing
    fig.update_layout(
        width=1100,
        height=500,
        autosize=True,
        margin=dict(t=110, b=100, l=100, r=250),
        title="Sinistros por Gênero e Gravidade (2019 a 2022)<br>Rodovia: " + Rodovia
    )

    fig.update_layout(legend=dict(
        orientation="v",
        y=1.03,
        x=1.02
    ))

    fig.update_layout(paper_bgcolor="rgba(230, 230, 230, 0.0)", plot_bgcolor="rgba(230, 230, 230, 0.3)")

    hoje = datetime.today().strftime('%d-%m-%Y')
    caminho = '16_' + str(Trecho) + '_box_sex_year_' + str(hoje) + '.png'
    pio.write_image(fig, "images/" + caminho)

    document.add_picture("images/" + caminho, height=Mm(7 * 10), width=Mm(15.42 * 10))

    st.plotly_chart(fig)
def box_sex():
    fig = px.box(df_pessoas2, x='Tipo Pessoa Envolvida', y='Idade Pessoa',
                 color="Sexo Pessoa", facet_row='Ano Ocorrencia', facet_col='classificacao_acidente',
                 color_discrete_map={"M": '#050df7', 'F': '#f70505'},
                 category_orders={'Tipo Pessoa Envolvida': ['Condutor', 'Passageiro'],
                                  'Sexo Pessoa': ['M', 'F'],
                                  'Ano Ocorrencia': ['2019', '2020', '2021', '2022'],
                                  'classificacao_acidente': ['Sem Vítimas', 'Com Vítimas Feridas',
                                                             'Com Vítimas Fatais']},
                 points='all',
                 labels={'Sexo Pessoa': 'Gênero',
                         'Tipo Pessoa Envolvida': '',
                         'Idade Pessoa': 'Idade'}
                 )

    fig.update_layout(uniformtext_minsize=8, uniformtext_mode='hide')

    fig.for_each_annotation(lambda a: a.update(text=a.text.split("=")[1]))

    fig.update_yaxes(
        #        title_text = "Sinistros",
        title_standoff=25,
        dtick=10,
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        range=[0, df_pessoas2['Idade Pessoa'].max() + 20])

    fig.update_xaxes(
        title_standoff=25,
        dtick=1,
        #        griddash='dash',
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        #        range=[data['Mês Ocorrencia'].min()-0.5, data['Nº Km'].max()+0.5]
    )

    # Update plot sizing
    fig.update_layout(
        width=1300,
        height=1300,
        autosize=True,
        margin=dict(t=100, b=100, l=100, r=250),
        title="Sinistros por Ano, Gênero e Gravidade<br>Rodovia: " + Rodovia
    )

    fig.update_layout(legend=dict(
        orientation="v",
        y=1.01,
        x=1.02
    ))

    fig.update_layout(paper_bgcolor="rgba(230, 230, 230, 0.0)", plot_bgcolor="rgba(230, 230, 230, 0.3)")

    hoje = datetime.today().strftime('%d-%m-%Y')
    caminho = '17_' + str(Trecho) + '_box_sex_' + str(hoje) + '.png'
    pio.write_image(fig, "images/" + caminho)

    document.add_picture("images/" + caminho, height=Mm(15.23 * 10), width=Mm(15.23 * 10))

    st.plotly_chart(fig)
def box_class():
    fig = px.box(df_pessoas2, x='Tipo Pessoa Envolvida', y='Idade Pessoa',
                 color="classificacao_acidente",
                 color_discrete_map={"Sem Vítimas": 'green', 'Com Vítimas Feridas': '#ff7f0e',
                                     'Com Vítimas Fatais': '#cc0202'},
                 category_orders={'Tipo Pessoa Envolvida': ['Condutor', 'Passageiro'],
                                  'Sexo Pessoa': ['M', 'F'],
                                  'classificacao_acidente': ['Sem Vítimas', 'Com Vítimas Feridas',
                                                             'Com Vítimas Fatais']},
                 points='all',
                 labels={'classificacao_acidente': 'Classificação do Sinistro',
                         'Tipo Pessoa Envolvida': '',
                         'Idade Pessoa': 'Idade'}
                 )

    fig.update_layout(uniformtext_minsize=8, uniformtext_mode='hide')

    fig.for_each_annotation(lambda a: a.update(text=a.text.split("=")[1]))

    fig.update_yaxes(
        #        title_text = "Sinistros",
        title_standoff=25,
        dtick=10,
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        range=[0, df_pessoas2['Idade Pessoa'].max() + 20])

    fig.update_xaxes(
        title_standoff=25,
        dtick=1,
        #        griddash='dash',
        showline=True,
        linecolor='black',
        mirror=True,
        showgrid=True,
        gridcolor='LightGrey',
        #        range=[data['Mês Ocorrencia'].min()-0.5, data['Nº Km'].max()+0.5]
    )

    # Update plot sizing
    fig.update_layout(
        width=1000,
        height=500,
        autosize=True,
        margin=dict(t=100, b=100, l=100, r=250),
        title="Sinistros por Ano, Gênero e Gravidade<br>Rodovia: " + Rodovia
    )

    fig.update_layout(legend=dict(
        orientation="v",
        y=1.02,
        x=1.02
    ))

    fig.update_layout(paper_bgcolor="rgba(230, 230, 230, 0.0)", plot_bgcolor="rgba(230, 230, 230, 0.3)")

    hoje = datetime.today().strftime('%d-%m-%Y')
    caminho = '18_' + str(Trecho) + '_box_class_' + str(hoje) + '.png'
    pio.write_image(fig, "images/" + caminho)

    document.add_picture("images/" + caminho, height=Mm(8 * 10), width=Mm(16 * 10))

    st.plotly_chart(fig)
def generate_word_report():
    hoje = datetime.today().strftime('%d-%m-%Y')
    document.save('documento/Relatório' + str(Trecho) + '_' + str(Rodovia) + '_' + str(hoje) + '.docx')
    st.write("Word report generated!")

# Cria um valor compartilhado entre as threads
progress = Value('i', 0)

def sidebar():
    # Criar campos para entrada de informações na sidebar
    rodovia = st.sidebar.text_input("Rodovia:")
    trecho = st.sidebar.text_input("Trecho:")
    ano = st.sidebar.text_input("Ano:")
    data_base = st.sidebar.text_input("Data Base:")
    km_inicial = st.sidebar.text_input("Km Inicial:")
    km_final = st.sidebar.text_input("Km Final:")

    # Criar botão de atualizar
    if st.sidebar.button("Atualizar"):
        sinistros_ano()

# Informações para o usuário inserir

Rodovia = chart_visual_rodovia

Trecho = "todos"

Ano = chart_visual_ano

DataBase = '02/01/2023'

km_inicial = 10
km_final = 150

Velocidade = 100
Classe = "-"
Situacao = "Pavimentada"

TMDA = 10769
TMDAAno = "2022"

Intervalo_conf = 1.645
intervalo = "95%"

Taxa_de_correção = 3.5

IA_medio = 1.17
NA_medio = 1.09

IA_medio_class_total = 1.29
NA_medio_class_total = 0.52

IA_medio_pav_total = 1.08
NA_medio_pav_total = 0.93

# Informações que chegam da planilha

Num_acidentes = 58

Extensao = 31.12

Num_dias = 365

Num_vitimas = 57

Num_acid_s_vitimas = 23

Num_acid_c_vitimas = 32

Num_acid_c_obitos = 3

Num_feridos = 54

Num_obitos = 3

# Índice Geral

# Massa de exposição
Massa_exposicao = (TMDA * Num_dias * Extensao) / (10 ** 6)

# Unidade Padrão de Severidade
UPS = 1 * Num_acid_s_vitimas + 5 * Num_acid_c_vitimas + 13 * Num_acid_c_obitos

# Indice para Redes

# Número Anual de Acidentes por KM
NA = (Num_acidentes * 365) / (Extensao * Num_dias)

# Índice de Acidentes
IA = (Num_acidentes) / ((Massa_exposicao))

# Índice Relacionados às Consequências

# Índice de Periculosidade
IP = Num_vitimas / (Num_acid_c_vitimas + Num_acid_c_obitos)

# Índice de Gravidade

if Num_obitos > 0:
    IG = Num_feridos / Num_obitos
else:
    IG = '-'

    # Índice de Severidade
IS = (UPS) / (Massa_exposicao)

# Índice de Controle

# Geral

# Número crítico de acidentes por Km do Trecho
NC = (NA_medio) + (Intervalo_conf * (math.sqrt(NA_medio / Extensao))) + (1 / (2 * Extensao))

# Índice Crítico
IC = (IA_medio) + (Intervalo_conf * (math.sqrt(IA_medio / Massa_exposicao))) + (1 / (2 * Massa_exposicao))

# Tipo de Pavimento

# Número crítico de acidentes por Km do Trecho
NC_pav = (NA_medio_pav_total) + (Intervalo_conf * (math.sqrt(NA_medio_pav_total / Extensao))) + (1 / (2 * Extensao))

# Índice Crítico
IC_pav = (IA_medio_pav_total) + (Intervalo_conf * (math.sqrt(IA_medio_pav_total / Massa_exposicao))) + (
        1 / (2 * Massa_exposicao))

# Tipo de Classe

# Número crítico de acidentes por Km do Trecho
NC_class = (NA_medio_class_total) + (Intervalo_conf * (math.sqrt(NA_medio_class_total / Extensao))) + (
        1 / (2 * Extensao))

# Índice Crítico
IC_class = (IA_medio_class_total) + (Intervalo_conf * (math.sqrt(IA_medio_class_total / Massa_exposicao))) + (
        1 / (2 * Massa_exposicao))

# Calculo da Criticidade

if NC < NA:
    CriticoNA = "Critico por Numero Anual de Acidentes"
else:
    CriticoNA = "Não Critico por Numero Anual de Acidentes"

if IC < IA:
    CriticoIA = "Critico por Índice de Acidentes"
else:
    CriticoIA = "Não Critico por Índice de Acidentes"

data_referencia = [['Rodovia', Rodovia],
                   ['Trecho', Trecho],
                   ['Segmento', (str(km_inicial) + ' ao ' + str(km_final))],
                   ['da Classe', Classe],
                   ['do Pavimento', Situacao],
                   ['Extensão', Extensao],
                   ['Velocidade Diretriz', Velocidade],
                   ['Ano dos Sinistros', Ano],
                   ['Data Base', DataBase]]

col_names = ['Entrada', 'Atributos']

st.table(data_referencia)

table = document.add_table(rows=1, cols=len(col_names))
table.style = 'Light List Accent 1'

headers = table.rows[0].cells

for i, name in enumerate(col_names):
    headers[i].text = name

for row in data_referencia:
    cells = table.add_row().cells
    for i, item in enumerate(row):
        cells[i].text = str(item)

data_indices = [['TMDA', TMDA],
                ['Ano do levantamento do TMDA', TMDAAno],
                #             ['Taxa de Correção',Taxa_de_correcao],
                ['Intervalo de Confiança', intervalo],
                ['--------------------------------------------', '-----------'],
                ['Índices Gerais:', ''],
                ['', ''],
                ['Unidade Padrão de Severidade (UPS)', UPS],
                ['Massa de Exposição (ME)', Massa_exposicao],
                ['--------------------------------------------', '-----------'],
                ['Índice de Controle:', ''],
                ['', ''],
                ['Média:', ''],
                ['Número Anual de Acidentes por KM (NC):', NC],
                ['Índice de Acidentes (IC):', IC],
                ['', ''],
                ['Média por pavimento:', ''],
                ['Número Anual de Acidentes por KM (NC):', NC_pav],
                ['Índice de Acidentes (IC):', IC_pav],
                ['', ''],
                ['Média por classe:', ''],
                ['Número Anual de Acidentes por KM (NC):', NC_class],
                ['Índice de Acidentes (IC):', IC_class],
                ['--------------------------------------------', '-----------'],
                ['Índices para Redes:', ''],
                ['', ''],
                ['Número Anual de Acidentes por KM (NA):', NA],
                ['Índice de Acidentes (IA):', IA],
                ['--------------------------------------------', '-----------'],
                ['Índice Relacionados às Consequencias:', ''],
                ['', ''],
                ['Indice de Periculosidade (IP):', IP],
                ['Índice de Gravidade (IG):', IG],
                ['Índice de Severidade (IS):', IS],
                ['', ''],
                ['', ''],
                ['O trecho é:', ""],
                [CriticoIA, ""],
                [CriticoNA, ""]]

col_names = ['Indices', 'Atributos']

st.table(data_indices)

table2 = document.add_table(rows=1, cols=len(col_names))
table2.style = 'Light List Accent 1'

headers = table2.rows[0].cells

for i, name in enumerate(col_names):
    headers[i].text = name

for row in data_indices:
    cells = table2.add_row().cells
    for i, item in enumerate(row):
        cells[i].text = str(item)

# Tabela para tratamento

# Tratamento para quantidade de Sinistros

# Retirada de Protocolos duplicados
df_ocorrencias = df.drop_duplicates(subset="Protocolo BATEU")

# Retirada de Ocorrencias duplicadas por KM/Rodovia/Dia
df_ocorrencias = df_ocorrencias.drop_duplicates(subset=("Nº Km", 'Rodovia', 'Data Fato', 'Veiculo N°'), keep="first")

# Retirada de Ocorrencias de Municipios que não obedecem a regra criada
df_ocorrencias = df_ocorrencias.loc[df_ocorrencias['municipio_considerar'] == 1]

# Retirada de Ocorrencias em Rodovias que não fazem parte da jurisdição do DER
df_ocorrencias = df_ocorrencias.loc[df_ocorrencias['principal'] == 1]

####XXXXXPossivel Criar Variavel especifica para cada Definição posteriormenteXXXXX####

# Definição da Rodovia de referência
df_ocorrencias = df_ocorrencias.loc[df_ocorrencias['Rodovia'] == Rodovia]

# Definição do Trecho de referência
   #df_ocorrencias = df_ocorrencias.loc[df_ocorrencias['trecho'] == Trecho]

# Definição do intervalo de interesse
df_ocorrencias = df_ocorrencias.loc[(df_ocorrencias['Nº Km'] >= km_inicial) & (df_ocorrencias['Nº Km'] <= km_final)]

# Tratamento para quantidade de pessoas envolvidas nos Sinistros

# Retirada de Ocorrencias duplicadas por Protocolo/Veiculo/Pessoa
df_pessoas = df.drop_duplicates(subset=("Protocolo BATEU", 'Veiculo N°', "Pessoa N°"), keep="first")

# Retirada de Ocorrencias de Municipios que não obedecem a regra criada
df_pessoas = df_pessoas.loc[df_pessoas['municipio_considerar'] == 1]

# Retirada de Ocorrencias em Rodovias que não fazem parte da jurisdição do DER
df_pessoas = df_pessoas.loc[df_pessoas['principal'] == 1]

# Definição da Rodovia de referência
df_pessoas = df_pessoas.loc[df_pessoas['Rodovia'] == Rodovia]

# Definição do Trecho de referência
# df_pessoas = df_pessoas.loc[df_pessoas['trecho'] == Trecho]

# Definição do intervalo de interesse
df_pessoas = df_pessoas.loc[(df_pessoas['Nº Km'] >= km_inicial) & (df_pessoas['Nº Km'] <= km_final)]

# Tratamento para quantidade de veiculos envolvidas nos Sinistros

# Retirada de Ocorrencias duplicadas por Protocolo/Veiculo/Pessoa
df_veiculos = df.drop_duplicates(subset=("Protocolo BATEU", 'Veiculo N°'), keep="first")

# Retirada de Ocorrencias de Municipios que não obedecem a regra criada
df_veiculos = df_veiculos.loc[df_veiculos['municipio_considerar'] == 1]

# Retirada de Ocorrencias em Rodovias que não fazem parte da jurisdição do DER
df_veiculos = df_veiculos.loc[df_veiculos['principal'] == 1]

# Definição da Rodovia de referência
df_veiculos = df_veiculos.loc[df_veiculos['Rodovia'] == Rodovia]

# Definição do Trecho de referência
# df_veiculos = df_veiculos.loc[df_veiculos['trecho'] == Trecho]

# Definição do intervalo de interesse
df_veiculos = df_veiculos.loc[(df_veiculos['Nº Km'] >= km_inicial) & (df_pessoas['Nº Km'] <= km_final)]

# df_pessoas['Resultado Vítima'] = df_pessoas['Resultado Vítima'].astype('string')
df_pessoas['Resultado Vítima'] = df_pessoas['Resultado Vítima'].fillna('Não Informado')
# df_pessoas['Resultado Vítima'] = df_pessoas['Resultado Vítima'].astype('category')

df_class_pessoas_ano = df_pessoas.pivot_table(index='Protocolo BATEU', columns='Resultado Vítima', values='Protocolo',
                                              aggfunc='count').reset_index()

df_class_pessoas_ano['Ilesos'] = df_class_pessoas_ano['Sem Ferimentos'] + df_class_pessoas_ano['Não Informado']
df_class_pessoas_ano.drop(['Não Informado', 'Sem Ferimentos'], axis=1, inplace=True)
df_class_pessoas_ano['Óbitos'] = df_class_pessoas_ano['Óbito no Local'] + df_class_pessoas_ano['Óbito Posterior']
df_class_pessoas_ano.drop(['Óbito no Local', 'Óbito Posterior'], axis=1, inplace=True)
df_class_pessoas_ano = df_class_pessoas_ano.rename(columns={'Ferimento': 'Feridos'})
df_class_pessoas_ano['classificacao_acidente'] = np.nan

for i in range(len(df_class_pessoas_ano)):

    if df_class_pessoas_ano['Óbitos'][i] > 0:
        df_class_pessoas_ano['classificacao_acidente'][i] = 2

    elif df_class_pessoas_ano['Feridos'][i] > 0:
        df_class_pessoas_ano['classificacao_acidente'][i] = 1

    else:
        df_class_pessoas_ano['classificacao_acidente'][i] = 0

df_class_pessoas_ano['classificacao_acidente'] = df_class_pessoas_ano['classificacao_acidente'].map(
    {2: 'Com Vítimas Fatais',
     1: 'Com Vítimas Feridas',
     0: 'Sem Vítimas'}).astype(str)

df_acidentes_ano = df_ocorrencias.merge(df_class_pessoas_ano, how='left', on=['Protocolo BATEU'])

data = df_acidentes_ano.groupby(
        ['Ano Ocorrencia', 'dia_da_semana', 'classificacao_acidente', 'Hora Ocorrencia', 'Rodovia'])[
        'Protocolo BATEU'].count().reset_index().rename(columns={'Protocolo BATEU': 'Acidentes'})

data = df_acidentes_ano.groupby(['Ano Ocorrencia', 'classificacao_acidente', 'Rodovia'])[
    'Protocolo BATEU'].count().reset_index().rename(columns={'Protocolo BATEU': 'Acidentes'})

sinistros_ano()

data = df_acidentes_ano.groupby(['Ano Ocorrencia', 'dia_da_semana', 'classificacao_acidente', 'Hora Ocorrencia', 'Rodovia'])[
    'Protocolo BATEU'].count().reset_index().rename(columns={'Protocolo BATEU': 'Acidentes'})

semana_hora()
semana_hora_year_class()
semana_hora_year()
semana_hora_class()

df_class_veiculos = df_pessoas.pivot_table(index=['Protocolo BATEU', 'Tipo Veículo'], columns='Resultado Vítima',
                                           values='Protocolo', aggfunc='count').reset_index()
df_class_veiculos['Ilesos'] = df_class_veiculos['Sem Ferimentos'] + df_class_veiculos['Não Informado']
df_class_veiculos.drop(['Não Informado', 'Sem Ferimentos'], axis=1, inplace=True)
df_class_veiculos['Óbitos'] = df_class_veiculos['Óbito no Local'] + df_class_veiculos['Óbito Posterior']
df_class_veiculos.drop(['Óbito no Local', 'Óbito Posterior'], axis=1, inplace=True)
df_class_veiculos = df_class_veiculos.rename(columns={'Ferimento': 'Feridos'})
df_class_veiculos['classificacao_acidente'] = np.nan

for i in range(len(df_class_veiculos)):

    if df_class_veiculos['Óbitos'][i] > 0:
        df_class_veiculos['classificacao_acidente'][i] = 2

    elif df_class_veiculos['Feridos'][i] > 0:
        df_class_veiculos['classificacao_acidente'][i] = 1

    elif df_class_veiculos['Ilesos'][i] > 0:
        df_class_veiculos['classificacao_acidente'][i] = 0

df_class_veiculos = df_class_veiculos.dropna()

df_class_veiculos['classificacao_acidente'] = df_class_veiculos['classificacao_acidente'].map({2: 'Com Vítimas Fatais',
                                                                                               1: 'Com Vítimas Feridas',
                                                                                               0: 'Sem Vítimas'}).astype(
    str)

df_acidentes_veiculos = df_ocorrencias.merge(df_class_veiculos, how='right', on=['Protocolo BATEU'])

df_acidentes_veiculos['Tipo Veículo_x'] = df_acidentes_veiculos['Tipo Veículo_x'].map({'AUTOMOVEL': 'AUTOMÓVEL',
                                                                                       'BICICLETA': 'BICICLETA',
                                                                                       'CAMINHAO': 'CAMINHÃO',
                                                                                       'CAMINHAO TRATOR': 'CAMINHÃO',
                                                                                       'REBOQUE': 'CAMINHÃO',
                                                                                       'SEMI-REBOQUE': 'CAMINHÃO',
                                                                                       'CAMINHONETE': 'UTILITÁRIO',
                                                                                       'CAMIONETA': 'UTILITÁRIO',
                                                                                       'UTILITARIO': 'UTILITÁRIO',
                                                                                       'CARRINHO DE PUXAR MATERIAIS RECICLAVEIS': 'OUTROS',
                                                                                       'CARROCA': 'OUTROS',
                                                                                       'CHARRETE': 'OUTROS',
                                                                                       'MOTOR CASA': 'OUTROS',
                                                                                       'NÃO INFORMADO': 'OUTROS',
                                                                                       'OUTROS': 'OUTROS',
                                                                                       'PATINS': 'OUTROS',
                                                                                       'TRATOR DE RODAS': 'OUTROS',
                                                                                       'TRATOR MISTO': 'OUTROS',
                                                                                       'TREM': 'OUTROS',
                                                                                       'CICLOMOTOR': 'MOTOCICLETA',
                                                                                       'MOTOCICLETA': 'MOTOCICLETA',
                                                                                       'MOTONETA': 'MOTOCICLETA',
                                                                                       'TRICICLO': 'MOTOCICLETA',
                                                                                       'MICRO ONIBUS': 'ONIBUS',
                                                                                       'ONIBUS': 'ONIBUS'
                                                                                       })

df_veiculos = df_acidentes_veiculos

df_veiculos1 = df_veiculos.groupby(['Ano Ocorrencia', 'Tipo Veículo_x', 'Rodovia', 'classificacao_acidente'])[
    'Protocolo BATEU'].count().reset_index().rename(columns={'Protocolo BATEU': 'Acidentes'})

veiculo_year()
veiculo()
km_year()
km()
mes_year()
mes()
tipo_year()
tipo()
clima_year()
clima()

df_pessoas['Sexo Pessoa'] = df_pessoas['Sexo Pessoa'].cat.remove_categories([' ', 'N'])

df_pessoas2 = df_pessoas.merge(df_class_pessoas_ano, how='left', on=['Protocolo BATEU'])

df_pessoas2.dropna(subset='Sexo Pessoa', inplace=True)

# Filtrar o dataframe original para incluir somente as linhas com valor diferente de zero na coluna desejada
df_pessoas2 = df_pessoas2[df_pessoas2['Idade Pessoa'] > 0]

box_sex_year()

box_sex()

box_class()

if st.sidebar.button("Generate Word Report"):
    generate_word_report()
